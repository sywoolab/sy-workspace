"""셀트리온 IM 마크다운 → Word 변환"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

INPUT = "/Users/sywoo/Library/CloudStorage/OneDrive-개인/바탕 화면/workspace/_data/260313_셀트리온_report_final.md"
OUTPUT = "/Users/sywoo/Library/CloudStorage/OneDrive-개인/바탕 화면/workspace/_data/260313_셀트리온_IM.docx"

doc = Document()

# 페이지 설정 A4
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# 스타일 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

# 한글 폰트 설정
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = '맑은 고딕'
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    h_style.font.color.rgb = RGBColor(0, 51, 102)
    if level == 1:
        h_style.font.size = Pt(16)
    elif level == 2:
        h_style.font.size = Pt(13)
    else:
        h_style.font.size = Pt(11)

with open(INPUT, 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')
i = 0

def add_table_from_lines(doc, table_lines):
    """마크다운 테이블을 Word 테이블로 변환"""
    rows = []
    for line in table_lines:
        line = line.strip()
        if line.startswith('|') and not re.match(r'^\|[\s\-:]+\|', line):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < num_cols:
                cell = table.cell(ri, ci)
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles['Normal']
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = '맑은 고딕'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    if ri == 0:
                        for run in paragraph.runs:
                            run.bold = True

    doc.add_paragraph()

def process_inline(text):
    """볼드, 이탤릭 등 인라인 마크업 처리"""
    text = text.strip()
    # 제거할 마크다운 기호
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # 볼드는 텍스트만 유지
    text = re.sub(r'\*(.+?)\*', r'\1', text)  # 이탤릭
    return text

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # 빈 줄
    if not stripped:
        i += 1
        continue

    # 구분선
    if stripped == '---':
        i += 1
        continue

    # 코드 블록
    if stripped.startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        if code_lines:
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            run = p.add_run('\n'.join(code_lines))
            run.font.size = Pt(9)
            run.font.name = 'Consolas'
        i += 1
        continue

    # 제목
    if stripped.startswith('# ') and not stripped.startswith('## '):
        text = stripped[2:]
        doc.add_heading(process_inline(text), level=1)
        i += 1
        continue

    if stripped.startswith('## '):
        text = stripped[3:]
        doc.add_heading(process_inline(text), level=2)
        i += 1
        continue

    if stripped.startswith('### '):
        text = stripped[4:]
        doc.add_heading(process_inline(text), level=3)
        i += 1
        continue

    # 테이블
    if stripped.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        add_table_from_lines(doc, table_lines)
        continue

    # 인용
    if stripped.startswith('> '):
        text = process_inline(stripped[2:])
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        pf = p.paragraph_format
        pf.left_indent = Cm(1)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)
        run.italic = True
        i += 1
        continue

    # 리스트
    if stripped.startswith('- ') or stripped.startswith('* '):
        text = process_inline(stripped[2:])
        p = doc.add_paragraph(text, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        i += 1
        continue

    # 숫자 리스트
    if re.match(r'^\d+\.\s', stripped):
        text = re.sub(r'^\d+\.\s', '', stripped)
        text = process_inline(text)
        p = doc.add_paragraph(text, style='List Number')
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        i += 1
        continue

    # 일반 텍스트
    text = process_inline(stripped)
    if text:
        p = doc.add_paragraph(text)
        p.style = doc.styles['Normal']
        for run in p.runs:
            run.font.name = '맑은 고딕'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    i += 1

doc.save(OUTPUT)
print(f"저장 완료: {OUTPUT}")
