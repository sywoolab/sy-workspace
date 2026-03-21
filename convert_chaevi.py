"""채비 심사신청서 별첨 마크다운 → Word 변환"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT = "/Users/sywoo/Library/CloudStorage/OneDrive-개인/바탕 화면/workspace/_data/260316_채비_심사신청서_v2.md"
OUTPUT = "/Users/sywoo/Library/CloudStorage/OneDrive-개인/바탕 화면/workspace/_data/260316_채비_심사신청서_v2.docx"

doc = Document()

# 페이지 설정 A4
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# 스타일 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = '맑은 고딕'
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    h_style.font.color.rgb = RGBColor(0, 51, 102)
    if level == 1:
        h_style.font.size = Pt(16)
        h_style.paragraph_format.space_before = Pt(24)
        h_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h_style.font.size = Pt(13)
        h_style.paragraph_format.space_before = Pt(18)
        h_style.paragraph_format.space_after = Pt(8)
    else:
        h_style.font.size = Pt(11)
        h_style.paragraph_format.space_before = Pt(12)
        h_style.paragraph_format.space_after = Pt(6)

# Heading 4 스타일 추가 (####)
try:
    h4_style = doc.styles['Heading 4']
except KeyError:
    h4_style = doc.styles.add_style('Heading 4', 1)  # WD_STYLE_TYPE.PARAGRAPH
h4_style.font.name = '맑은 고딕'
h4_style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
h4_style.font.color.rgb = RGBColor(0, 51, 102)
h4_style.font.size = Pt(10.5)
h4_style.font.bold = True
h4_style.paragraph_format.space_before = Pt(10)
h4_style.paragraph_format.space_after = Pt(4)

with open(INPUT, 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')
i = 0


def set_cell_shading(cell, color_hex):
    """셀 배경색 설정"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_table_from_lines(doc, table_lines):
    """마크다운 테이블을 Word 테이블로 변환"""
    rows = []
    for line in table_lines:
        line = line.strip()
        if line.startswith('|') and not re.match(r'^\|[\s\-:]+\|', line):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 테이블 너비 설정 (전체 페이지 너비)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < num_cols:
                cell = table.cell(ri, ci)
                # 인라인 마크업 제거
                cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                cell_text = re.sub(r'\*(.+?)\*', r'\1', cell_text)
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles['Normal']
                    paragraph.paragraph_format.space_after = Pt(1)
                    paragraph.paragraph_format.space_before = Pt(1)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = '맑은 고딕'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    if ri == 0:
                        set_cell_shading(cell, '003366')
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()


def process_inline(text):
    """볼드, 이탤릭 등 인라인 마크업 처리"""
    text = text.strip()
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    return text


def add_paragraph_with_bold(doc, text):
    """볼드 마크업을 실제 볼드 서식으로 변환하여 추가"""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']

    # **text** 패턴을 찾아서 볼드 처리
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            # 이탤릭도 처리
            sub_parts = re.split(r'(\*.+?\*)', part)
            for sub in sub_parts:
                if sub.startswith('*') and sub.endswith('*') and not sub.startswith('**'):
                    run = p.add_run(sub[1:-1])
                    run.italic = True
                else:
                    run = p.add_run(sub)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(10)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return p


while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # 빈 줄
    if not stripped:
        i += 1
        continue

    # 구분선
    if stripped == '---':
        # 페이지 나누기 대신 공백
        doc.add_paragraph()
        i += 1
        continue

    # 코드 블록
    if stripped.startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        if code_lines:
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            pf = p.paragraph_format
            pf.left_indent = Cm(0.5)
            run = p.add_run('\n'.join(code_lines))
            run.font.size = Pt(9)
            run.font.name = 'Consolas'
        i += 1
        continue

    # 제목 (#### 먼저 체크)
    if stripped.startswith('#### '):
        text = stripped[5:]
        doc.add_heading(process_inline(text), level=4)
        i += 1
        continue

    if stripped.startswith('### '):
        text = stripped[4:]
        doc.add_heading(process_inline(text), level=3)
        i += 1
        continue

    if stripped.startswith('## '):
        text = stripped[3:]
        doc.add_heading(process_inline(text), level=2)
        i += 1
        continue

    if stripped.startswith('# '):
        text = stripped[2:]
        doc.add_heading(process_inline(text), level=1)
        i += 1
        continue

    # 테이블
    if stripped.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        add_table_from_lines(doc, table_lines)
        continue

    # 인용 (출처 표기 등)
    if stripped.startswith('> '):
        text = process_inline(stripped[2:])
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        pf = p.paragraph_format
        pf.left_indent = Cm(1)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)
        run.italic = True
        i += 1
        continue

    # (출처: ...) 단독 줄
    if stripped.startswith('(출처:'):
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run(process_inline(stripped))
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        i += 1
        continue

    # 리스트
    if stripped.startswith('- ') or stripped.startswith('* '):
        text = stripped[2:]
        p = doc.add_paragraph(style='List Bullet')
        # 볼드 마크업 처리
        parts = re.split(r'(\*\*.+?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        i += 1
        continue

    # 숫자 리스트
    if re.match(r'^\d+\.\s', stripped):
        text = re.sub(r'^\d+\.\s', '', stripped)
        text = process_inline(text)
        p = doc.add_paragraph(text, style='List Number')
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        i += 1
        continue

    # 일반 텍스트 (볼드 마크업 포함 가능)
    if '**' in stripped:
        add_paragraph_with_bold(doc, stripped)
    else:
        text = process_inline(stripped)
        if text:
            p = doc.add_paragraph(text)
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.font.name = '맑은 고딕'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    i += 1

doc.save(OUTPUT)
print(f"저장 완료: {OUTPUT}")
